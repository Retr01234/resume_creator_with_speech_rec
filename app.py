from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
  pyttsx3.speak(text)

document = Document()

# Adding Image
document.add_picture(
  'pic.jpeg', 
  width = Inches(2.0),
  height = Inches(2.0)
)

# Main Contact Field
speak('What is Your Name?')

name = input('Name: ')
speak('Welcome ' + name + '. I am Roby and I will create you a Resume')

speak('What is your Phone Number?')
number = input('Phone Number: ')

speak('What is Your E-Mail?')
email = input('E-Mail: ')

document.add_paragraph(
  name + ' | ' + number + ' | ' + email,
)

speak('Thank You ' + name)

# About Us Section
document.add_heading('About Me:')
speak('Please tell me about yourself')
document.add_paragraph(
  input('About You: ')
)

speak('Thank You')

# Work Experience
document.add_heading('Work Experience')
speak('Now tell me about your previous work experience')
p = document.add_paragraph()

speak('What is the Name of the Company?')
company = input('Company: ')

speak('What day did you start?')
start_date = input('Start Date: ')

speak('What day did you stop working there?')
end_date = input('End Date: ')

speak('What was your position?')
position = input('Position: ')

# Adding Text to Existing Paragraphs
p.add_run(company + ' ' + '\n').bold = True
p.add_run(start_date + '-' + end_date + '\n').italic = True
p.add_run(position + ' ' + '\n').bold = True

speak('What was your role at ' + company)
experience_details = input(
  'What is your Role at ' + company + '? '
)

p.add_run(experience_details)

# Adding More Work Experience
while True:
  speak('Do you have any more experience? Type yes to tell us about it, or type no if not')
  has_more_experiences = input('Do you have any more experience? Yes or No: ')

  if has_more_experiences.lower() == 'yes':
    p = document.add_paragraph()


    speak('What is the Name of the Company?')
    company = input('Company: ')

    speak('What day did you start?')
    start_date = input('Start Date: ')

    speak('What day did you stop working there?')
    end_date = input('End Date: ')

    speak('What was your position?')
    position = input('Position: ')

    p.add_run(company + ' ' + '\n').bold = True
    p.add_run(start_date + '-' + end_date + '\n').italic = True
    p.add_run(position + ' ' + '\n')

    speak('What was your role at ' + company)
    experience_details = input(
      'What is your Role at ' + company + '? '
    )

    p.add_run(experience_details)
  
  else:
    break

# Adding Skillset
speak('Now to wrap this up. Tell me about your Skill set')
document.add_heading('IT Skills:')
p = document.add_paragraph()

speak('What Skill do you have?')
skill = input('Skill: ')

speak('How much experience do you have with it?')
amount_of_exp = input('Experience: ')

p.add_run('\u2022 ' + skill + ':').bold = True
p.add_run(amount_of_exp)

# More SKills
while True:
  speak('Do you have any more skills? Type yes, if you do. And type no, if not.')
  has_more_skills = input('Any more skills? Yes or No: ')

  if has_more_skills.lower() == 'yes':
    p = document.add_paragraph()

    speak('What Skill do you have?')
    skill = input('Skill: ')

    speak('How much experience do you have?')
    amount_of_exp = input('Experience: ')

    p.add_run('\u2022 ' + skill + ': ').bold = True
    p.add_run(amount_of_exp)
  
  else:
    break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "Resume Generated with my great Skills"

speak('Thank You for your Input. I will go ahead and create your Resume. Goodbye.')

document.save('Resume.docx') # Saves All Changes to the Word File